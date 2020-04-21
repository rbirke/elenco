#!/bin/env python

from docx import Document
import sqlite3
import re
import pickle

LEVELS = ['I', "II", 'III', 'IV', 'V', 'VI', 'VII']

def split_spheres(spheres):
    spheres = spheres.upper().split(',')
    for i, sphere in enumerate(spheres):
        csphere = sphere.lstrip().rstrip()
        spheres[i] = csphere
        words = csphere.split()
        if 'E' in words:
            j = words.index('E')
            spheres[i] = " ".join(words[0:j])
            if words[j+1] in ['ARIA','ACQUA','FUOCO','TERRA']:                    
                expanded_sphere = [words[0]]
                expanded_sphere.extend(words[j+1:])
            else:
                expanded_sphere = []
                expanded_sphere.extend(words[j+1:])
            spheres.append(" ".join(expanded_sphere))
    spheres.sort()
    return spheres


def extract_gods(filename):
    document = Document(filename)

    gods = {}

    for p in document.paragraphs:
        if 'DIVINITÀ' in p.text.upper().split()[0:2]:
            group = p.text.upper().lstrip().rstrip()
            gods[group] = {}
        else:
            t = p.text.split(':')
            if len(t) == 2:
                if t[0].lower() == 'sfere maggiori':
                    gods[group][god]['major'] = split_spheres(t[1])
                elif t[0].lower() == 'sfere minori':
                    gods[group][god]['minor'] = split_spheres(t[1])
            elif len(p.runs) > 1 and p.runs[0].bold == True:
                god = p.runs[0].text.lstrip().rstrip()
                gods[group][god] = {'name' : god}

    return gods

def extract_spells(filename):
    document = Document(filename)

    # Extract spheres
    spheres = []

    for p in document.paragraphs:
        sphere = p.text.lstrip().rstrip()
        if sphere != '' and sphere not in LEVELS and len(sphere.split()) < 3:
            spheres.append(sphere.lower())

    #Extract spells        
    level = 0
    i = 0
    spells = {spheres[i].upper(): []}
    for table in document.tables:
        spell = {}
        for row in table.rows:
            for cell in row.cells:
                key_value = re.split('[:?]', cell.text)
                key = key_value[0].lstrip().rstrip()
                value = ":".join(key_value[1:]).lstrip().rstrip()

                if key.lower() == 'livello':
                    # Convert roman number
                    if value.upper() in LEVELS:
                        value = LEVELS.index(value.upper()) + 1
                    else:
                        # Remove non numeric characters
                        value = re.sub('[^0-9]', '', value)    
                spell[key] = value
        

        l = int(spell['Livello'])
        if l < level: # Next sphere
            i += 1
            spells[spheres[i].upper()] = []
            level = 0

        if l > level: # Next spell level in the same sphere
            spells[spheres[i].upper()].append([])
            level = l

        spell['Sfera'] = spheres[i].upper()
        spells[spheres[i].upper()][level-1].append(spell)
    return spells


def get_form():
    from ipywidgets import VBox, GridBox, Checkbox, Label, Layout, Dropdown, Text, HTML, Button, IntProgress
    from docx.shared import Pt
    from docx.enum.style import WD_STYLE_TYPE
    import operator
    from base64 import b64encode

    spells = pickle.load( open( "spells.p", "rb" ) )
    gods = pickle.load( open( "gods.p", "rb" ) )

    sphere_checkboxes = { 'major': {}, 'minor': {}}

    def on_change(b):
        html.value = ''
        progress.value = 0

    for sphere in spells.keys():
        for t in ['major', 'minor']:
            sphere_checkboxes[t][sphere] = Checkbox(value=False, description=sphere, indent=False)
            sphere_checkboxes[t][sphere].observe(on_change, names='value')

    gods_per_race = Dropdown(
        options=list(gods.keys()),
        description='Gruppi di divinità:',
        value=None
    )

    gods_of_race = Dropdown(
        options=[],
        description='Divinità:',
    )

    extra_spheres = {
        'major': Text(value='',
        placeholder='Sfere non specificate appaiono qui',
        description='Sfere maggiori',
        disabled=True),
        'minor' : Text(value='',
        placeholder='Sfere non specificate appaiono qui',
        description='Sfere minori',
        disabled=True)
    }

    def on_change_per_race(change):
        html.value = ''
        progress.value = 0
        v = change['new']
        gods_of_race.value = None
        if v is None:
            gods_of_race.options = None
        else:
            gods_of_race.options = list(zip(gods[v].keys(), gods[v].values()))
    gods_per_race.observe(on_change_per_race, names='value')

    def on_change_per_god(change):
        html.value = ''
        progress.value = 0
        for t in extra_spheres:
            extra_spheres[t].value = ''
        god = change['new']
        if god is None:
            for t in ['major', 'minor']:
                for c in sphere_checkboxes[t]:
                    sphere_checkboxes[t][c].value = False
        else:
            for t in sphere_checkboxes:
                if t in god:
                    for c in sphere_checkboxes[t]:
                        if c in god[t]:
                            sphere_checkboxes[t][c].value = True
                        else:
                            sphere_checkboxes[t][c].value = False
                    for s in god[t]:
                        if s not in sphere_checkboxes[t]:
                            if extra_spheres[t].value is None:
                                extra_spheres[t].value = s
                            else:
                                extra_spheres[t].value += s
                            
                else:
                    for c in sphere_checkboxes[t]:
                        sphere_checkboxes[t][c].value = False
    gods_of_race.observe(on_change_per_god, names='value')

    button = Button(
        description='Genera elenco',
        disabled=False
    )

    def add_content(document, content, space_after=0, font_name='Arial', font_size=16, line_spacing=0, space_before=0, keep_together=True, keep_with_next=False, page_break_before=False,
                    widow_control=False, set_bold=False, set_italic=False, set_underline=False, set_all_caps=False,style_name=""):
        paragraph = document.add_paragraph(content)
        font = paragraph.style.font
        font.name = font_name
        font.size = Pt(font_size)
        font.bold = set_bold
        font.italic = set_italic
        font.all_caps = set_all_caps
        font.underline = set_underline
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = Pt(space_before)
        paragraph_format.space_after = Pt(space_after)
        paragraph_format.line_spacing = line_spacing
        paragraph_format.keep_together = keep_together
        paragraph_format.keep_with_next = keep_with_next
        paragraph_format.page_break_before = page_break_before
        paragraph_format.widow_control = widow_control

    def build_list(b):
        progress.max = 7 * len(sphere_checkboxes['major']) + 3 * len(sphere_checkboxes['minor'])
                            
        document = Document()
        style = document.styles.add_style('Indent', WD_STYLE_TYPE.PARAGRAPH)
        paragraph_format = style.paragraph_format
        paragraph_format.space_before = Pt(3)
        paragraph_format.widow_control = True
        
        title = 'Lista Incantesimi'
        if gods_of_race.value is not None and 'name' in gods_of_race.value:
            title = gods_of_race.value['name']
        document.add_heading(title, 0)

        for i in range(7):
            spells_level = []
            for c in sphere_checkboxes['major']:
                progress.value += 1
                if sphere_checkboxes['major'][c].value == True:
                    lst = spells[c][i]
                    for s in lst:
                        s['Sfera'] = s['Sfera'].upper()
                    spells_level.extend(lst)
            if i < 3:
                for c in sphere_checkboxes['minor']:
                    progress.value += 1
                    if sphere_checkboxes['minor'][c].value == True:
                        lst = spells[c][i]
                        for s in lst:
                            s['Sfera'] = s['Sfera'].lower()
                        spells_level.extend(lst)
            # Sort by name
            spells_level.sort(key=operator.itemgetter('Incantesimo', 'Effetto'))
            
            if len(spells_level) > 0:
                document.add_heading('Incantesimi di livello '+str(i+1), 1)
            
            def key_value(cell, key, value, font_size=None):
                p = cell.paragraphs[0]
                if font_size is not None:
                    p.style.font.size = font_size
                r = p.add_run()
                r.text = key
                r.font.bold = True
                r = p.add_run()
                r.text = ': ' + str(value)
                
            spells_level_len = len(spells_level)
            i = 0
            while i < spells_level_len:
                # Separator
                paragraph = document.add_paragraph()
                paragraph.style.font.size = Pt(3)
                paragraph_format = paragraph.paragraph_format
                paragraph_format.space_before = Pt(0)
                paragraph_format.space_after = Pt(0)

                # Count and skip duplicates
                c = [spells_level[i]['Sfera']]
                while i+1 < spells_level_len and \
                    spells_level[i]['Incantesimo'] == spells_level[i+1]['Incantesimo'] and \
                    spells_level[i]['Effetto'] == spells_level[i+1]['Effetto']:
                    c.append(spells_level[i+1]['Sfera'])
                    i += 1
                
                s = spells_level[i]

                # Spell in table format
                table = document.add_table(4, 4)
                table.style = 'TableGrid'
                cells_0 = table.rows[0].cells
                key_value(cells_0[0], 'Incantesimo', s['Incantesimo'], Pt(11))
                cells_0[0].merge(cells_0[1])
                key_value(cells_0[2], '#', str(len(c)))
                cells_0[3].text = ', '.join(c)

                cells_1 = table.rows[1].cells
                key_value(cells_1[0], 'Livello', s['Livello'])
                key_value(cells_1[1], 'Raggio d’azione', s['Raggio d’azione'])
                key_value(cells_1[2], 'Componenti', s['Componenti'])
                key_value(cells_1[3], 'Durata', s['Durata'])

                cells_2 = table.rows[2].cells
                key_value(cells_2[0], 'Tempo lancio', s['Tempo lancio'])
                key_value(cells_2[1], 'Area d’effetto', s['Area d’effetto'])
                key_value(cells_2[2], 'Tiro salvezza', s['Tiro salvezza'])
                key_value(cells_2[3], 'Reversibile', s['Reversibile'])

                cells_3 = table.rows[3].cells
                key_value(cells_3[0], 'Effetto', s['Effetto'])
                cells_3[0].merge(cells_3[1])
                cells_3[0].merge(cells_3[2])
                cells_3[0].merge(cells_3[3])
                i += 1
                
        # Save and generate downlpoad link
        document.save('lista_incantesimi.docx')
        with open('lista_incantesimi.docx', 'rb') as fp:
            data = str(b64encode(fp.read()))
            html_ = "<a download='{filename}' href='data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;charset=utf-8;base64,{payload}' target='_blank'>{title}</a>"
            html.value = html_.format(payload=data[2:], title="Scarica Lista Incantesimi", filename=title+'.docx')
                

    button.on_click(build_list)
    html = HTML(value="")


    progress = IntProgress(
        value=0,
        min=0,
        max=10,
        step=1,
        description='',
        bar_style='', # 'success', 'info', 'warning', 'danger' or ''
        orientation='horizontal',
        layout=Layout(width='95%')
    )

    form = VBox([
        GridBox([gods_per_race, gods_of_race], layout=Layout(grid_template_columns="repeat(2, 450px)")),
        GridBox(list(extra_spheres.values()), layout=Layout(grid_template_columns="repeat(2, 450px)")),
        Label(value='Sfere maggiori:'),
        GridBox(list(sphere_checkboxes['major'].values()), layout=Layout(grid_template_columns="repeat(4, 200px)")),
        Label(value='Sfere minori:'),
        GridBox(list(sphere_checkboxes['minor'].values()), layout=Layout(grid_template_columns="repeat(4, 200px)")),
        button,
        progress,
        html
    ])
    return form

if __name__ == "__main__":
    gods = extract_gods('divinita.docx')    
    spells = extract_spells('magie.docx')
    # Save them as pickles
    pickle.dump( gods, open( "gods.p", "wb" ) )
    pickle.dump( spells, open( "spells.p", "wb" ) )